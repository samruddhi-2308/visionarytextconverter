from flask import Flask, request, render_template, send_file
import os
from PIL import Image
from spellchecker import SpellChecker
import pytesseract
from docx import Document
from deep_translator import GoogleTranslator
from transformers import pipeline
from transformers import BlipProcessor, BlipForConditionalGeneration

# Initialize the processor and model
processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
summarizer = pipeline("summarization")

pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

# Upload folder
app = Flask(__name__)
UPLOAD_FOLDER = 'upload'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload')
def upload_page():
    return render_template('upload.html')

@app.route('/faq_page')
def faq_page():
    return render_template('faq.html')

@app.route('/choose_page')
def choose_page():
    return render_template('choose.html')

@app.route('/generate_page')
def generate_page():
    return render_template('generate.html')

@app.route('/summarize_page', methods=['GET', 'POST'])
def summarize_page():
    if request.method == "POST":
        text = request.form.get("inputText")  # Use .get() to avoid crashing on missing keys

        if not text:  # Handle case where inputText is missing
            return render_template("summarize.html", result="No text provided to summarize!")

        try:
            summarized_text = summarizer(text, max_length=150, min_length=50, do_sample=False)[0]["summary_text"]
        except Exception as e:
            return render_template("summarize.html", result=f"Error summarizing text: {str(e)}")

        return render_template("summarize.html", result=summarized_text)

    return render_template("summarize.html", result=None)

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return "No file part in request", 400
    file = request.files['file']
    if file.filename == '':
        return "No file selected", 400
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(filepath)

    try:
        # Open the image
        image = Image.open(filepath)

        # Extract text using PyTesseract
        raw_text = pytesseract.image_to_string(image, lang='eng', config='--psm 6')

        # Perform spell checking
        words = raw_text.split()
        spell = SpellChecker()
        corrected_words = [spell.correction(word) if word.isalpha() else word for word in words]
        corrected_text = ' '.join(corrected_words)

        # Remove the uploaded file after processing
        os.remove(filepath)

        # Render the result
        return render_template('result.html', text=corrected_text)

    except Exception as e:
        os.remove(filepath)  # Ensure file is removed in case of an error
        return f"Error processing the image: {str(e)}", 500

@app.route('/save', methods=['POST'])
def save():
    text = request.form['text']
    doc = Document()
    doc.add_heading('Extracted Text', 0)
    doc.add_paragraph(text)
    word_filepath = 'extracted.docx'
    doc.save(word_filepath)

    return send_file(word_filepath, as_attachment=True, download_name='text.docx')

@app.route('/translate', methods=['POST'])
def translate():
    try:
        original_text = request.form['original_text']
        target_language = request.form['language']
        translated_text = GoogleTranslator(source='auto', target=target_language).translate(original_text)
        return render_template('result.html', text=original_text, translated_text=translated_text)
    except Exception as e:
        return f"Error during translation: {str(e)}", 500


@app.route('/generate_page', methods=['POST'])
def get_image_caption():
    if 'file' not in request.files:
        return "No file part in request", 400
    file = request.files['file']
    if file.filename == '':
        return "No file selected", 400
    # Process the image here
    raw_image = Image.open(file).convert('RGB')  # Assuming file is used here
    inputs = processor(raw_image, return_tensors="pt")
    out = model.generate(**inputs)
    generated_text = processor.decode(out[0], skip_special_tokens=True)
    
    # Pass the generated text to the template
    return render_template('generate.html', text=generated_text)

if __name__ == '__main__':
    app.run(debug=True)
